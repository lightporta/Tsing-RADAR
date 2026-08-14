"""A6 生成产物、扫描、签名交付与并发重放边界。"""

from __future__ import annotations

import hashlib
import hmac
import io
import json
import logging
import posixpath
import threading
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from pathlib import Path
from urllib.parse import urlsplit

from docx import Document
from fastapi import HTTPException
from fastapi.testclient import TestClient
from lxml import etree
from pypdf import PdfReader
import pytest

from app.core.config import settings
from app.core.logging_filters import ArtifactTokenRedactionFilter
from app.db.session import SessionLocal
from app.main import app, startup_event
from app.models.private_document import (
    ArtifactDeliveryGrant,
    DeletedArtifactTombstone,
    PrivateDocument,
)
from app.services.artifact_delivery import (
    issue_delivery_grant,
    redeem_delivery_token,
)
from app.services.artifact_generation import _find_cjk_font
from app.services.identity import Principal
from app.services.object_storage import ObjectStorageError
from app.services.private_documents import (
    delete_private_document_consistently,
    delete_private_document_file,
    read_private_document_bytes,
)

QXD_BEARER = "test-qxd-key"
QXD_CLAIM_SECRET = "test-qxd-end-user-secret"


def test_linux_auto_font_prefers_reportlab_compatible_wqy(monkeypatch):
    wqy = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"
    noto = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
    monkeypatch.setattr(settings, "DOCUMENT_CJK_FONT_PATH", None)
    monkeypatch.setattr(
        Path,
        "is_file",
        lambda path: str(path) in {wqy, noto},
    )
    assert _find_cjk_font() == Path(wqy)


def _web_client() -> tuple[TestClient, dict[str, str]]:
    client = TestClient(app)
    response = client.get("/api/session")
    assert response.status_code == 200
    return client, {"X-CSRF-Token": client.cookies["tsing_radar_csrf"]}


def _idem(
    headers: dict[str, str],
    key: str | None = None,
) -> dict[str, str]:
    return {
        **headers,
        "Idempotency-Key": key or f"test:{uuid.uuid4()}",
    }


def _minimal_docx(text: str = "A6 私有简历") -> bytes:
    output = io.BytesIO()
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:body></w:document>"
    )
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="xml" ContentType="application/xml"/>'
                "</Types>"
            ),
        )
        archive.writestr("word/document.xml", document)
    return output.getvalue()


def _upload_document(
    client: TestClient,
    headers: dict[str, str],
    *,
    name: str = "resume.docx",
) -> dict:
    response = client.post(
        "/api/documents",
        headers=headers,
        files={
            "file": (
                name,
                _minimal_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


def _confirmed_interview(
    client: TestClient,
    headers: dict[str, str],
) -> str:
    session_id = str(uuid.uuid4())
    response = client.post(
        "/api/v1/llm/chat",
        params={"stream": "false"},
        headers=headers,
        json={
            "session_id": session_id,
            "messages": [
                {"role": "user", "content": "自然语言处理、对话系统"},
                {"role": "user", "content": "工程落地"},
                {"role": "user", "content": "高频具体指导"},
                {"role": "user", "content": "产业就业"},
                {"role": "user", "content": "成熟稳妥路线"},
                {"role": "user", "content": "无"},
                {"role": "user", "content": "确认画像"},
            ],
        },
    )
    assert response.status_code == 200, response.text
    assert response.json()["recommend_ready"] is True
    return session_id


def _qxd_headers(claim: str | None = None) -> dict[str, str]:
    headers = {"Authorization": f"Bearer {QXD_BEARER}"}
    if claim is not None:
        headers["X-QXD-End-User-Id"] = claim
        headers["X-QXD-End-User-Signature"] = hmac.new(
            QXD_CLAIM_SECRET.encode(),
            claim.encode(),
            hashlib.sha256,
        ).hexdigest()
    return headers


def _token_from_url(url: str) -> str:
    return urlsplit(url).path.rsplit("/", 1)[-1]


def _grant_id(token: str) -> str:
    return token.split(".")[1]


def _assert_docx_package_integrity(payload: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        assert archive.testzip() is None
        names = set(archive.namelist())
        assert {
            "[Content_Types].xml",
            "_rels/.rels",
            "word/document.xml",
        }.issubset(names)
        assert "docProps/custom.xml" not in names

        for name in names:
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            root = etree.fromstring(archive.read(name))
            for node in root.iter():
                ignorable = node.attrib.get(
                    "{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable"
                )
                if ignorable:
                    assert set(ignorable.split()).issubset(node.nsmap)

            if not name.endswith(".rels"):
                continue
            if name == "_rels/.rels":
                source_dir = ""
            else:
                prefix, rel_file = name.rsplit("/_rels/", 1)
                source_dir = posixpath.dirname(
                    posixpath.join(prefix, rel_file.removesuffix(".rels"))
                )
            for relation in root:
                if relation.attrib.get("TargetMode") == "External":
                    continue
                target = posixpath.normpath(
                    posixpath.join(source_dir, relation.attrib["Target"])
                ).lstrip("/")
                assert target in names, (name, target)

        content_types = etree.fromstring(archive.read("[Content_Types].xml"))
        for override in content_types:
            part_name = override.attrib.get("PartName")
            if part_name:
                assert part_name.lstrip("/") in names

        story_xml = b"".join(
            archive.read(name)
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        assert b"rsid" not in story_xml


def test_upload_records_honest_scan_scope_and_rejects_eicar():
    client, headers = _web_client()
    item = _upload_document(client, headers)
    assert item["status"] == "ready"
    assert item["scan_status"] == "clean"
    assert item["scan_scope"] == "structural_signature_only"
    assert not {"stored_name", "object_key", "path", "url"} & set(item)

    eicar = (
        b"%PDF-1.4\n"
        b"X5O!P%@AP[4\\PZX54(P^)7CC)7}$"
        b"EICAR-STANDARD-ANTIVIRUS-TEST-FILE!$H+H*"
    )
    rejected = client.post(
        "/api/documents",
        headers=headers,
        files={"file": ("eicar.pdf", eicar, "application/pdf")},
    )
    assert rejected.status_code == 422
    assert "反病毒测试特征" in rejected.json()["detail"]


def test_resume_pdf_docx_generation_requires_confirmation_and_private_download():
    owner, owner_headers = _web_client()
    other, _other_headers = _web_client()
    base = {
        "student_name": "测试同学",
        "dept": "自动化系",
        "email": "student@example.edu",
        "phone": "13800000000",
        "education": "本科三年级",
        "research_interests": ["自然语言处理", "对话系统"],
        "projects": [{"name": "课程项目", "detail": "实现可复现实验"}],
        "awards": ["校级奖学金"],
        "positions": ["学生社团项目负责人"],
        "format": "pdf",
    }
    refused = owner.post(
        "/api/resume/generate",
        headers=_idem(owner_headers),
        json={**base, "confirm_generation": False},
    )
    assert refused.status_code == 422

    generated = owner.post(
        "/api/resume/generate",
        headers=_idem(owner_headers),
        json={**base, "confirm_generation": True},
    )
    assert generated.status_code == 200, generated.text
    pdf_item = generated.json()
    assert pdf_item["document_kind"] == "resume"
    assert pdf_item["media_type"] == "application/pdf"
    assert pdf_item["scan_status"] == "clean"

    denied = owner.post(
        f"/api/artifacts/{pdf_item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": False},
    )
    assert denied.status_code == 422
    issued = owner.post(
        f"/api/artifacts/{pdf_item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": True},
    )
    assert issued.status_code == 200
    download_url = issued.json()["download_url"]
    token = _token_from_url(download_url)
    with SessionLocal() as db:
        grant = db.get(ArtifactDeliveryGrant, _grant_id(token))
        assert grant is not None
        assert token not in grant.token_digest
        assert pdf_item["document_id"] not in token

    # 越权不会消费额度；原会话随后仍可成功一次。
    assert other.post(download_url, headers=_other_headers).status_code == 403
    downloaded = owner.post(download_url, headers=owner_headers)
    assert downloaded.status_code == 200
    assert downloaded.content.startswith(b"%PDF-")
    assert "filename*=UTF-8''" in downloaded.headers["content-disposition"]
    assert owner.post(download_url, headers=owner_headers).status_code == 410
    text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(io.BytesIO(downloaded.content)).pages
    )
    assert "测试同学" in text
    assert "系统未核验经历真实性" in text

    docx_generated = owner.post(
        "/api/resume/generate",
        headers=_idem(owner_headers),
        json={**base, "format": "docx", "confirm_generation": True},
    )
    assert docx_generated.status_code == 200, docx_generated.text
    docx_item = docx_generated.json()
    grant = owner.post(
        f"/api/artifacts/{docx_item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": True},
    ).json()
    content = owner.post(grant["download_url"], headers=owner_headers).content
    document = Document(io.BytesIO(content))
    assert document.core_properties.author in {"", None}
    assert document.core_properties.last_modified_by in {"", None}
    assert "测试同学" in "\n".join(p.text for p in document.paragraphs)
    _assert_docx_package_integrity(content)
    section = document.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11
    assert round(section.top_margin.inches, 2) == 1
    assert round(section.right_margin.inches, 2) == 1
    assert round(section.bottom_margin.inches, 2) == 1
    assert round(section.left_margin.inches, 2) == 1


def test_match_report_is_real_download_and_honest_when_published_count_is_zero():
    client, headers = _web_client()
    session_id = _confirmed_interview(client, headers)
    response = client.post(
        "/api/artifacts/match-report",
        headers=_idem(headers),
        json={
            "session_id": session_id,
            "format": "pdf",
            "confirm_generation": True,
        },
    )
    assert response.status_code == 200, response.text
    item = response.json()
    assert item["document_kind"] == "match_report"
    grant = client.post(
        f"/api/artifacts/{item['document_id']}/download-grant",
        headers=_idem(headers),
        json={"confirm_private_download": True},
    ).json()
    downloaded = client.post(grant["download_url"], headers=headers)
    assert downloaded.status_code == 200
    text = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(io.BytesIO(downloaded.content)).pages
    )
    assert "暂无通过审核的数据" in text
    assert "本报告不包含导师推荐名单" in text
    assert "2027 官方招生目录" in text
    assert "推荐你" not in text

    docx_response = client.post(
        "/api/artifacts/match-report",
        headers=_idem(headers),
        json={
            "session_id": session_id,
            "format": "docx",
            "confirm_generation": True,
        },
    )
    assert docx_response.status_code == 200
    docx_grant = client.post(
        f"/api/artifacts/{docx_response.json()['document_id']}/download-grant",
        headers=_idem(headers),
        json={"confirm_private_download": True},
    ).json()
    docx_payload = client.post(
        docx_grant["download_url"],
        headers=headers,
    ).content
    _assert_docx_package_integrity(docx_payload)
    docx_document = Document(io.BytesIO(docx_payload))
    docx_text = "\n".join(
        [paragraph.text for paragraph in docx_document.paragraphs]
        + [
            cell.text
            for table in docx_document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "暂无通过审核的数据" in docx_text
    assert "本报告不包含导师推荐名单" in docx_text


def test_invalid_tampered_expired_and_wrong_owner_tokens_do_not_consume():
    owner, owner_headers = _web_client()
    other, other_headers = _web_client()
    item = _upload_document(owner, owner_headers)
    issued = owner.post(
        f"/api/artifacts/{item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": True},
    ).json()
    token = _token_from_url(issued["download_url"])
    grant_id = _grant_id(token)

    tampered = token.replace(token.split(".")[2], "known-low-entropy-value")
    assert (
        owner.post(
            f"/api/artifacts/download/{tampered}",
            headers=owner_headers,
        ).status_code
        == 404
    )
    assert other.post(issued["download_url"], headers=other_headers).status_code == 403
    with SessionLocal() as db:
        assert db.get(ArtifactDeliveryGrant, grant_id).use_count == 0

    with SessionLocal() as db:
        grant = db.get(ArtifactDeliveryGrant, grant_id)
        grant.expires_at = datetime.now(timezone.utc) - timedelta(seconds=1)
        db.commit()
    assert owner.post(issued["download_url"], headers=owner_headers).status_code == 410
    with SessionLocal() as db:
        assert db.get(ArtifactDeliveryGrant, grant_id).use_count == 0


def _race_redeem(
    token: str,
    audience: str,
    principal: Principal | None,
    barrier: threading.Barrier,
) -> int:
    barrier.wait()
    with SessionLocal() as db:
        try:
            redeem_delivery_token(
                db,
                token=token,
                audience=audience,
                principal=principal,
            )
            return 200
        except HTTPException as exc:
            return exc.status_code


def test_atomic_web_and_qxd_consumption_under_concurrency():
    client, headers = _web_client()
    item = _upload_document(client, headers)
    with SessionLocal() as db:
        document = db.get(PrivateDocument, item["document_id"])
        owner_subject_id = document.owner_subject_id
        web_principal = Principal(
            subject_id=owner_subject_id,
            channel="web",
            auth_session_id="test",
            persistent=True,
        )
        web_issued = issue_delivery_grant(
            db,
            document=document,
            principal=web_principal,
            audience="web_private",
            confirmed=True,
            idempotency_key=f"direct-web:{uuid.uuid4()}",
        )
    web_token = _token_from_url(web_issued.download_url)
    web_barrier = threading.Barrier(2)
    with ThreadPoolExecutor(max_workers=2) as pool:
        web_results = list(
            pool.map(
                lambda _index: _race_redeem(
                    web_token,
                    "web_private",
                    web_principal,
                    web_barrier,
                ),
                range(2),
            )
        )
    assert web_results.count(200) == 1
    with SessionLocal() as db:
        web_grant = db.get(ArtifactDeliveryGrant, _grant_id(web_token))
        assert web_grant.use_count == 1
        assert web_grant.revoked is True

    with SessionLocal() as db:
        document = db.get(PrivateDocument, item["document_id"])
        document.document_kind = "match_report"
        db.commit()
        qxd_principal = Principal(
            subject_id=owner_subject_id,
            channel="qxd",
            auth_session_id=None,
            persistent=True,
        )
        qxd_issued = issue_delivery_grant(
            db,
            document=document,
            principal=qxd_principal,
            audience="qxd_platform",
            confirmed=True,
            idempotency_key=f"direct-qxd:{uuid.uuid4()}",
        )
    qxd_token = _token_from_url(qxd_issued.download_url)
    qxd_barrier = threading.Barrier(6)
    with ThreadPoolExecutor(max_workers=6) as pool:
        qxd_results = list(
            pool.map(
                lambda _index: _race_redeem(
                    qxd_token,
                    "qxd_platform",
                    None,
                    qxd_barrier,
                ),
                range(6),
            )
        )
    assert qxd_results.count(200) == 3
    with SessionLocal() as db:
        qxd_grant = db.get(ArtifactDeliveryGrant, _grant_id(qxd_token))
        assert qxd_grant.use_count == 3
        assert qxd_grant.revoked is True


def test_missing_object_burns_consumption_without_storage_detail():
    client, headers = _web_client()
    item = _upload_document(client, headers)
    issued = client.post(
        f"/api/artifacts/{item['document_id']}/download-grant",
        headers=_idem(headers),
        json={"confirm_private_download": True},
    ).json()
    token = _token_from_url(issued["download_url"])
    with SessionLocal() as db:
        document = db.get(PrivateDocument, item["document_id"])
        delete_private_document_file(document)
    response = client.post(issued["download_url"], headers=headers)
    assert response.status_code == 410
    assert response.json()["detail"] == "文件当前不可交付"
    with SessionLocal() as db:
        grant = db.get(ArtifactDeliveryGrant, _grant_id(token))
        assert grant.use_count == 1
        assert grant.revoked is True


def test_same_length_object_tampering_is_rejected_by_web_and_qxd_downloads():
    owner, owner_headers = _web_client()

    web_item = _upload_document(owner, owner_headers, name="web-resume.docx")
    with SessionLocal() as db:
        web_document = db.get(PrivateDocument, web_item["document_id"])
        web_path = (
            Path(settings.object_storage_local_root) / web_document.stored_name
        )
        original_size = web_document.size_bytes
    web_path.write_bytes(b"X" * original_size)
    web_grant = owner.post(
        f"/api/artifacts/{web_item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": True},
    )
    assert web_grant.status_code == 200
    web_download = owner.post(
        web_grant.json()["download_url"],
        headers=owner_headers,
    )
    assert web_download.status_code == 410
    assert web_download.content != b"X" * original_size

    qxd_item = _upload_document(owner, owner_headers, name="qxd-report.docx")
    with SessionLocal() as db:
        qxd_document = db.get(PrivateDocument, qxd_item["document_id"])
        qxd_document.document_kind = "match_report"
        qxd_path = (
            Path(settings.object_storage_local_root) / qxd_document.stored_name
        )
        qxd_size = qxd_document.size_bytes
        owner_subject_id = qxd_document.owner_subject_id
        db.commit()
        qxd_grant = issue_delivery_grant(
            db,
            document=qxd_document,
            principal=Principal(
                subject_id=owner_subject_id,
                channel="qxd",
                auth_session_id=None,
                persistent=True,
            ),
            audience="qxd_platform",
            confirmed=True,
            idempotency_key=f"tamper-qxd:{uuid.uuid4()}",
        )
    qxd_path.write_bytes(b"Y" * qxd_size)
    qxd_download = owner.get(urlsplit(qxd_grant.download_url).path)
    assert qxd_download.status_code == 410
    assert qxd_download.content != b"Y" * qxd_size


def test_delete_revokes_grants_is_idempotent_and_preserves_owner_authorization():
    owner, owner_headers = _web_client()
    other, other_headers = _web_client()
    item = _upload_document(owner, owner_headers)
    issued = owner.post(
        f"/api/artifacts/{item['document_id']}/download-grant",
        headers=_idem(owner_headers),
        json={"confirm_private_download": True},
    ).json()
    grant_id = _grant_id(_token_from_url(issued["download_url"]))

    denied = other.request(
        "DELETE",
        f"/api/documents/{item['document_id']}",
        headers=_idem(other_headers),
        json={"confirm_delete": True},
    )
    assert denied.status_code == 403

    deleted = owner.request(
        "DELETE",
        f"/api/documents/{item['document_id']}",
        headers=_idem(owner_headers),
        json={"confirm_delete": True},
    )
    assert deleted.status_code == 200
    assert deleted.json() == {"status": "deleted", "idempotent": False}
    assert owner.post(issued["download_url"], headers=owner_headers).status_code == 410
    with SessionLocal() as db:
        assert db.get(ArtifactDeliveryGrant, grant_id) is None
        assert db.get(PrivateDocument, item["document_id"]) is None
        tombstone = db.get(DeletedArtifactTombstone, item["document_id"])
        assert tombstone is not None

    repeated = owner.request(
        "DELETE",
        f"/api/documents/{item['document_id']}",
        headers=_idem(owner_headers),
        json={"confirm_delete": True},
    )
    assert repeated.status_code == 200
    assert repeated.json() == {"status": "deleted", "idempotent": True}
    assert (
        other.request(
            "DELETE",
            f"/api/documents/{item['document_id']}",
            headers=_idem(other_headers),
            json={"confirm_delete": True},
        ).status_code
        == 403
    )


def test_object_delete_failure_keeps_non_deliverable_metadata_and_retry_converges(
    monkeypatch,
):
    client, headers = _web_client()
    item = _upload_document(client, headers)
    issued = client.post(
        f"/api/artifacts/{item['document_id']}/download-grant",
        headers=_idem(headers),
        json={"confirm_private_download": True},
    ).json()
    grant_id = _grant_id(_token_from_url(issued["download_url"]))

    import app.services.private_documents as private_document_service

    original_delete = private_document_service.delete_private_document_file

    def fail_delete(_document):
        raise ObjectStorageError("simulated storage outage")

    monkeypatch.setattr(
        private_document_service,
        "delete_private_document_file",
        fail_delete,
    )
    failed = client.request(
        "DELETE",
        f"/api/documents/{item['document_id']}",
        headers=_idem(headers),
        json={"confirm_delete": True},
    )
    assert failed.status_code == 503
    assert failed.json()["detail"] == "对象删除失败，文件已停止交付，可安全重试"
    with SessionLocal() as db:
        persisted = db.get(PrivateDocument, item["document_id"])
        assert persisted is not None
        assert persisted.status == "delete_failed"
        assert db.get(ArtifactDeliveryGrant, grant_id) is None
    assert client.post(issued["download_url"], headers=headers).status_code == 410

    monkeypatch.setattr(
        private_document_service,
        "delete_private_document_file",
        original_delete,
    )
    retried = client.request(
        "DELETE",
        f"/api/documents/{item['document_id']}",
        headers=_idem(headers),
        json={"confirm_delete": True},
    )
    assert retried.status_code == 200
    with SessionLocal() as db:
        assert db.get(PrivateDocument, item["document_id"]) is None
        assert db.get(DeletedArtifactTombstone, item["document_id"]) is not None


def test_final_metadata_commit_failure_remains_retryable_and_converges(monkeypatch):
    client, headers = _web_client()
    item = _upload_document(client, headers)

    with SessionLocal() as db:
        document = db.get(PrivateDocument, item["document_id"])
        original_commit = db.commit
        commit_calls = 0

        def fail_second_commit():
            nonlocal commit_calls
            commit_calls += 1
            if commit_calls == 2:
                raise RuntimeError("simulated final commit failure")
            original_commit()

        monkeypatch.setattr(db, "commit", fail_second_commit)
        with pytest.raises(HTTPException) as exc_info:
            delete_private_document_consistently(db, document=document)
        assert exc_info.value.status_code == 503
        assert exc_info.value.detail == "对象已删除，元数据清理待安全重试"
        monkeypatch.setattr(db, "commit", original_commit)

        persisted = db.get(PrivateDocument, item["document_id"])
        assert persisted is not None
        assert persisted.status == "deleting"
        with pytest.raises(ObjectStorageError):
            read_private_document_bytes(persisted)

        delete_private_document_consistently(db, document=persisted)
        assert db.get(PrivateDocument, item["document_id"]) is None
        assert db.get(DeletedArtifactTombstone, item["document_id"]) is not None


def test_qxd_real_report_confirmation_emits_attachment_once_in_stop_frame():
    claim = f"a6-qxd-{uuid.uuid4()}"
    client = TestClient(app)
    user_turns = [
        "自然语言处理、对话系统",
        "工程落地",
        "高频具体指导",
        "产业就业",
        "成熟稳妥路线",
        "无",
        "确认画像",
        "生成匹配报告",
    ]
    disclosure = client.post(
        "/v1/chat/completions",
        headers=_qxd_headers(claim),
        json={
            "user": "a6-report",
            "messages": [
                {"role": "user", "content": content}
                for content in user_turns
            ],
            "stream": False,
        },
    )
    assert disclosure.status_code == 200, disclosure.text
    assert "确认生成并通过清小搭附件交付" in disclosure.json()["choices"][0]["message"]["content"]
    assert "x_soda" not in disclosure.json()

    user_turns.append("确认生成并通过清小搭附件交付")
    generated = client.post(
        "/v1/chat/completions",
        headers=_qxd_headers(claim),
        json={
            "user": "a6-report",
            "messages": [
                {"role": "user", "content": content}
                for content in user_turns
            ],
            "stream": True,
        },
    )
    assert generated.status_code == 200, generated.text
    data_lines = [
        line.removeprefix("data:").strip()
        for line in generated.text.splitlines()
        if line.startswith("data:")
    ]
    assert data_lines[-1] == "[DONE]"
    frames = [json.loads(line) for line in data_lines[:-1]]
    assert all("x_soda" not in frame for frame in frames[:-1])
    attachment = frames[-1]["x_soda"]["attachments"][0]
    assert attachment["fileType"] == "pdf"
    assert attachment["fileUrl"].startswith(
        "https://agent.example.edu/v1/attachments/"
    )
    assert attachment["expiresAt"]
    token = _token_from_url(attachment["fileUrl"])
    local_url = f"/v1/attachments/{token}"
    assert client.get(local_url).status_code == 200
    assert client.get(local_url).status_code == 200
    assert client.get(local_url).status_code == 200
    assert client.get(local_url).status_code == 410


def test_qxd_unverified_request_cannot_generate_public_attachment():
    client = TestClient(app)
    response = client.post(
        "/v1/chat/completions",
        headers=_qxd_headers(),
        json={
            "messages": [
                {"role": "user", "content": "自然语言处理"},
                {"role": "user", "content": "工程落地"},
                {"role": "user", "content": "高频具体指导"},
                {"role": "user", "content": "产业就业"},
                {"role": "user", "content": "成熟稳妥路线"},
                {"role": "user", "content": "无"},
                {"role": "user", "content": "确认画像"},
                {"role": "user", "content": "生成匹配报告"},
            ],
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert "可验证、稳定的终端用户身份" in payload["choices"][0]["message"]["content"]
    assert "x_soda" not in payload


def test_clamav_unavailable_fails_closed_without_storing(monkeypatch):
    client, headers = _web_client()
    monkeypatch.setattr(settings, "FILE_SCAN_MODE", "clamav")
    monkeypatch.setattr(settings, "CLAMAV_HOST", None)
    response = client.post(
        "/api/documents",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                _minimal_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 503
    assert "扫描服务不可用" in response.json()["detail"]


def test_artifact_bearer_tokens_are_redacted_from_application_access_logs():
    token = "v1.known-grant.known-nonce.known-signature"
    record = logging.LogRecord(
        name="uvicorn.access",
        level=logging.INFO,
        pathname=__file__,
        lineno=1,
        msg='%s - "%s %s HTTP/%s" %d',
        args=(
            "127.0.0.1:1234",
            "GET",
            f"/v1/attachments/{token}",
            "1.1",
            200,
        ),
        exc_info=None,
    )
    assert ArtifactTokenRedactionFilter().filter(record) is True
    rendered = record.getMessage()
    assert token not in rendered
    assert "/v1/attachments/[REDACTED]" in rendered


@pytest.mark.asyncio
async def test_production_artifact_configuration_fails_closed(monkeypatch):
    valid = {
        "DEBUG": False,
        "AUTO_CREATE_SCHEMA": False,
        "ADMIN_TOKEN": "m" * 32,
        "SESSION_HMAC_SECRET": "s" * 32,
        "WEB_COOKIE_SECURE": True,
        "ARTIFACT_SIGNING_SECRET": "a" * 32,
        "FILE_SCAN_MODE": "clamav",
        "CLAMAV_HOST": "clamav.internal",
        "OBJECT_STORE_BACKEND": "s3",
        "S3_BUCKET": "private-artifacts",
        "S3_ACCESS_KEY_ID": "test-access",
        "S3_SECRET_ACCESS_KEY": "test-secret",
        "S3_SERVER_SIDE_ENCRYPTION": "AES256",
        "QXD_API_KEY": "production-qxd-bearer-" + "b" * 32,
        "QXD_END_USER_SIGNING_SECRET": "q" * 32,
        "PUBLIC_BASE_URL": "https://agent.tsingradar.cn",
        "ALLOW_TEST_PUBLIC_BASE_URL": False,
    }
    invalid_cases = [
        (
            {"ARTIFACT_SIGNING_SECRET": "s" * 32},
            "不同密钥",
        ),
        ({"S3_BUCKET": None}, "私有 S3"),
        ({"S3_SERVER_SIDE_ENCRYPTION": "none"}, "S3 服务端加密"),
        (
            {"QXD_END_USER_SIGNING_SECRET": None},
            "QXD_END_USER_SIGNING_SECRET",
        ),
        (
            {"PUBLIC_BASE_URL": "https://agent.example.edu/path"},
            "公网 HTTPS",
        ),
    ]
    for overrides, expected in invalid_cases:
        with monkeypatch.context() as scoped:
            for name, value in {**valid, **overrides}.items():
                scoped.setattr(settings, name, value)
            with pytest.raises(RuntimeError, match=expected):
                await startup_event()
