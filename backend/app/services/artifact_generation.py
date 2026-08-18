"""A6 确定性 PDF/DOCX 简历与匹配报告生成。

生成器只使用用户已确认输入和治理后的匹配结果，不调用外部 LLM，也不补写
未经提供的经历、导师事实或推荐结论。
"""

from __future__ import annotations

import html
import io
import re
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from lxml import etree as ElementTree
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.private_document import PrivateDocument
from app.schemas.artifacts import ResumeArtifactRequest
from app.schemas.interview import StudentPortrait
from app.services.interview import confirmed_portrait, get_session
from app.services.match_application import (
    MatchApplicationOutcome,
    run_confirmed_match,
)
from app.services.artifact_audit import (
    add_artifact_event,
    commit_artifact_event,
    validation_reason,
)
from app.services.idempotency import (
    UNEXPECTED_FAILURE_DETAIL,
    begin_idempotency,
    complete_idempotency,
    fail_idempotency,
)
from app.services.private_documents import (
    discard_private_artifact_object,
    store_private_artifact,
    validate_and_extract_document,
)

_BLUE = colors.HexColor("#2E74B5")
_DARK_BLUE = colors.HexColor("#1F4D78")
_INK = colors.HexColor("#1F2937")
_MUTED = colors.HexColor("#667085")
_LIGHT_FILL = colors.HexColor("#F2F4F7")
_WARN_FILL = colors.HexColor("#FFF4E5")
_DOCX_BLUE = RGBColor(0x2E, 0x74, 0xB5)
_DOCX_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
_DOCX_MUTED = RGBColor(0x66, 0x70, 0x85)
_CJK_FONT_NAME = "TsingRadarCJK"

_PORTRAIT_LABELS = {
    "research_mode": {
        "theory": "理论与原理",
        "engineering": "工程与落地",
        "mixed": "两者结合",
        "undecided": "暂不确定",
    },
    "mentorship_style": {
        "high_guidance": "高频具体指导",
        "balanced": "平衡",
        "autonomous": "自主探索",
        "undecided": "暂不确定",
    },
    "career_orientation": {
        "academic": "学术深造",
        "industry": "产业就业",
        "national_mission": "国家任务",
        "mixed": "混合选择",
        "undecided": "暂不确定",
    },
    "innovation_risk": {
        "pioneering": "高风险新方向",
        "balanced": "平衡",
        "mature": "成熟路径",
        "undecided": "暂不确定",
    },
}
_CONSTRAINT_LABELS = {
    "location": "地点",
    "weekly_commitment_days": "每周投入天数",
    "degree_stage": "学历阶段",
    "language": "语言",
    "confidentiality": "保密要求",
    "graduation_arrangement": "毕业安排",
    "department": "院系",
    "research_topic": "研究主题",
    "advisor_id": "导师",
}
_OPERATOR_LABELS = {
    "equals": "必须等于",
    "one_of": "必须属于",
    "excludes": "必须排除",
    "contains": "必须包含",
    "minimum": "至少",
    "maximum": "至多",
}
_MATCH_STATUS_LABELS = {
    "matched": "已生成证据化候选",
    "no_published_data": "暂无通过审核的数据",
    "needs_clarification": "需要先澄清硬性条件",
    "profile_not_confirmed": "画像尚未确认",
}
_MATCH_METHOD_LABELS = {
    "evidence-matching-v1": "证据化多目标匹配 v1",
}
_RETRIEVAL_MODE_LABELS = {
    "deterministic_lexical_fallback": "确定性词法回退（非语义 embedding）",
    "hybrid_provider": "可插拔混合召回",
}


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _find_cjk_font() -> Path:
    candidates = [
        settings.DOCUMENT_CJK_FONT_PATH,
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyh.ttf",
        "C:/Windows/Fonts/simhei.ttf",
        # ReportLab can register the packaged WQY TrueType collection directly.
        # Some Noto CJK collections use outlines unsupported by the installed
        # ReportLab/fontTools combination, so keep them as later fallbacks.
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    raise HTTPException(
        status_code=503,
        detail="文档生成缺少可嵌入中文字体，请配置 DOCUMENT_CJK_FONT_PATH",
    )


def _ensure_pdf_font() -> str:
    if _CJK_FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        path = _find_cjk_font()
        kwargs = {"subfontIndex": 0} if path.suffix.lower() == ".ttc" else {}
        pdfmetrics.registerFont(TTFont(_CJK_FONT_NAME, str(path), **kwargs))
        pdfmetrics.registerFontFamily(
            _CJK_FONT_NAME,
            normal=_CJK_FONT_NAME,
            bold=_CJK_FONT_NAME,
            italic=_CJK_FONT_NAME,
            boldItalic=_CJK_FONT_NAME,
        )
    return _CJK_FONT_NAME


def _pdf_styles() -> dict[str, ParagraphStyle]:
    font = _ensure_pdf_font()
    return {
        "title": ParagraphStyle(
            "Title",
            fontName=font,
            fontSize=23,
            leading=30,
            textColor=_INK,
            alignment=TA_LEFT,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            fontName=font,
            fontSize=12,
            leading=18,
            textColor=_MUTED,
            spaceAfter=14,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "Heading1",
            fontName=font,
            fontSize=16,
            leading=22,
            textColor=_BLUE,
            spaceBefore=16,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "Heading2",
            fontName=font,
            fontSize=13,
            leading=19,
            textColor=_DARK_BLUE,
            spaceBefore=12,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "Body",
            fontName=font,
            fontSize=10.5,
            leading=15,
            textColor=_INK,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "Small",
            fontName=font,
            fontSize=8.5,
            leading=12,
            textColor=_MUTED,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "center": ParagraphStyle(
            "Center",
            fontName=font,
            fontSize=10,
            leading=15,
            alignment=TA_CENTER,
            textColor=_INK,
            wordWrap="CJK",
        ),
    }


def _pdf_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    return Paragraph(html.escape(str(text or "")).replace("\n", "<br/>"), style)


def _pdf_header_footer(canvas, document) -> None:
    font = _ensure_pdf_font()
    canvas.saveState()
    canvas.setTitle(document.title)
    canvas.setAuthor("Tsing-RADAR")
    canvas.setFont(font, 8)
    canvas.setFillColor(_MUTED)
    canvas.drawString(inch, 10.45 * inch, "Tsing-RADAR · 私有生成文档")
    canvas.drawRightString(
        7.5 * inch,
        0.48 * inch,
        f"第 {canvas.getPageNumber()} 页",
    )
    canvas.restoreState()


def _pdf_key_value_table(
    rows: list[tuple[str, str]],
    styles: dict[str, ParagraphStyle],
) -> Table:
    data = [
        [
            _pdf_paragraph(label, styles["small"]),
            _pdf_paragraph(value or "未填写", styles["body"]),
        ]
        for label, value in rows
    ]
    table = Table(data, colWidths=[1.35 * inch, 5.15 * inch], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), _LIGHT_FILL),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#E4E7EC")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _pdf_callout(
    text: str,
    styles: dict[str, ParagraphStyle],
    *,
    warning: bool = False,
) -> Table:
    table = Table(
        [[_pdf_paragraph(text, styles["body"])]],
        colWidths=[6.5 * inch],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), _WARN_FILL if warning else _LIGHT_FILL),
                ("BOX", (0, 0), (-1, -1), 0.7, _BLUE),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return table


def _portrait_rows(profile: StudentPortrait) -> list[tuple[str, str]]:
    constraints = [
        (
            f"{_CONSTRAINT_LABELS[item.field.value]}"
            f"{_OPERATOR_LABELS[item.operator.value]}"
            f"{' / '.join(item.value)}"
        )
        for item in profile.hard_constraints or []
    ]
    return [
        ("研究兴趣", "、".join(profile.research_interests) or "未填写"),
        (
            "研究方式",
            _PORTRAIT_LABELS["research_mode"].get(
                profile.research_mode or "undecided",
                "暂不确定",
            ),
        ),
        (
            "指导偏好",
            _PORTRAIT_LABELS["mentorship_style"].get(
                profile.mentorship_style or "undecided",
                "暂不确定",
            ),
        ),
        (
            "生涯方向",
            _PORTRAIT_LABELS["career_orientation"].get(
                profile.career_orientation or "undecided",
                "暂不确定",
            ),
        ),
        (
            "创新风险",
            _PORTRAIT_LABELS["innovation_risk"].get(
                profile.innovation_risk or "undecided",
                "暂不确定",
            ),
        ),
        ("硬性条件", "；".join(constraints) or "无已确认硬性条件"),
    ]


def _resume_plain_text(request: ResumeArtifactRequest) -> str:
    lines = [
        request.student_name,
        request.dept,
        request.education,
        request.email,
        request.phone,
        "研究兴趣：" + "、".join(request.research_interests),
    ]
    for project in request.projects:
        lines.append(f"项目：{project.name} {project.detail}".strip())
    lines.extend(f"奖项：{item}" for item in request.awards)
    lines.extend(f"经历：{item}" for item in request.positions)
    if request.target_advisor:
        lines.append(f"用户填写意向对象：{request.target_advisor}")
    return "\n".join(item for item in lines if item)


def _render_resume_pdf(request: ResumeArtifactRequest) -> bytes:
    styles = _pdf_styles()
    # 简历使用更紧凑但仍可读的节奏，避免把隐私提示孤立到空白次页。
    styles["subtitle"].spaceAfter = 10
    styles["h1"].spaceBefore = 10
    styles["h1"].spaceAfter = 5
    styles["h2"].spaceBefore = 7
    styles["h2"].spaceAfter = 3
    styles["body"].spaceAfter = 3
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=f"{request.student_name}-个人简历",
        author="Tsing-RADAR",
    )
    story = [
        _pdf_paragraph(request.student_name, styles["title"]),
        _pdf_paragraph("个人简历 · 用户确认信息的确定性排版版本", styles["subtitle"]),
        _pdf_key_value_table(
            [
                ("院系", request.dept),
                ("教育背景", request.education),
                ("邮箱", request.email),
                ("电话", request.phone),
            ],
            styles,
        ),
    ]
    if request.research_interests:
        story.extend(
            [
                _pdf_paragraph("研究兴趣", styles["h1"]),
                _pdf_paragraph("、".join(request.research_interests), styles["body"]),
            ]
        )
    if request.projects:
        story.append(_pdf_paragraph("项目经历", styles["h1"]))
        for project in request.projects:
            story.append(
                KeepTogether(
                    [
                        _pdf_paragraph(project.name, styles["h2"]),
                        _pdf_paragraph(project.detail or "未填写项目说明", styles["body"]),
                    ]
                )
            )
    if request.awards:
        story.extend(
            [
                _pdf_paragraph("奖项与荣誉", styles["h1"]),
                _pdf_key_value_table(
                    [(str(index), item) for index, item in enumerate(request.awards, 1)],
                    styles,
                ),
            ]
        )
    if request.positions:
        story.extend(
            [
                _pdf_paragraph("任职与实践", styles["h1"]),
                _pdf_key_value_table(
                    [(str(index), item) for index, item in enumerate(request.positions, 1)],
                    styles,
                ),
            ]
        )
    if request.target_advisor:
        story.extend(
            [
                _pdf_paragraph("意向说明", styles["h1"]),
                _pdf_callout(
                    f"用户填写的意向对象：{request.target_advisor}。该名称未由系统核验，"
                    "不代表对方正在招生或已同意接收材料。",
                    styles,
                    warning=True,
                ),
            ]
        )
    story.extend(
        [
            Spacer(1, 4),
            _pdf_callout(
                "隐私与真实性提示：本文档只排版用户本次明确提供的信息，"
                "系统未核验经历真实性，也未向任何导师、邮箱或第三方发送文件。"
                "请在每次对外使用前逐项复核。",
                styles,
                warning=True,
            ),
        ]
    )
    document.build(
        story,
        onFirstPage=_pdf_header_footer,
        onLaterPages=_pdf_header_footer,
    )
    return output.getvalue()


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_run_font(run, size=8, color=_DOCX_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=8, color=_DOCX_MUTED)


def _set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"),
        "Microsoft YaHei",
    )
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        new_grid.append(grid_col)
    table._tbl.replace(old_grid, new_grid)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def _configure_docx(*, compact: bool) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.05 if compact else 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, _DOCX_BLUE, 10 if compact else 16, 5 if compact else 8),
        ("Heading 2", 13, _DOCX_BLUE, 8 if compact else 12, 4 if compact else 6),
        ("Heading 3", 12, _DOCX_DARK_BLUE, 6 if compact else 8, 3 if compact else 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = ""
    header_run = header.add_run("Tsing-RADAR · 私有生成文档")
    _set_run_font(header_run, size=8, color=_DOCX_MUTED)
    _add_page_number(section.footer.paragraphs[0])
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = (
        "由 Tsing-RADAR 确定性模板生成；未调用外部模型。"
    )
    return document


def _docx_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    _set_run_font(run, size=23, bold=True)
    sub = document.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    sub_run = sub.add_run(subtitle)
    _set_run_font(sub_run, size=12, color=_DOCX_MUTED)


def _docx_key_value_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        label_run = cells[0].paragraphs[0].add_run(label)
        _set_run_font(label_run, size=10, color=_DOCX_DARK_BLUE, bold=True)
        _shade_cell(cells[0], "F2F4F7")
        cells[1].text = ""
        value_run = cells[1].paragraphs[0].add_run(value or "未填写")
        _set_run_font(value_run, size=10.5)
    _set_table_geometry(table, [1944, 7416])
    return table


def _docx_callout(document: Document, text: str, *, warning: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_run_font(run, size=10)
    _shade_cell(cell, "FFF4E5" if warning else "F2F4F7")
    _set_table_geometry(table, [9360])


def _scrub_docx_metadata(payload: bytes) -> bytes:
    """删除个人属性、custom properties 与 rsid 会话标记。"""
    source = io.BytesIO(payload)
    output = io.BytesIO()
    rsid_pattern = re.compile(r"^\{[^}]+\}rsid")
    with zipfile.ZipFile(source) as archive, zipfile.ZipFile(
        output,
        "w",
        zipfile.ZIP_DEFLATED,
    ) as rewritten:
        for info in archive.infolist():
            if info.filename == "docProps/custom.xml":
                continue
            data = archive.read(info.filename)
            if info.filename == "docProps/core.xml":
                root = ElementTree.fromstring(data)
                for name in (
                    "{http://purl.org/dc/elements/1.1/}creator",
                    "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy",
                ):
                    node = root.find(name)
                    if node is not None:
                        node.text = ""
                data = ElementTree.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                )
            elif info.filename.startswith("word/") and info.filename.endswith(".xml"):
                root = ElementTree.fromstring(data)
                for node in root.iter():
                    for attribute in list(node.attrib):
                        if rsid_pattern.match(attribute):
                            del node.attrib[attribute]
                for parent in root.iter():
                    for child in list(parent):
                        local_name = child.tag.rsplit("}", 1)[-1]
                        if local_name.startswith("rsid"):
                            parent.remove(child)
                data = ElementTree.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                )
            elif info.filename == "[Content_Types].xml":
                root = ElementTree.fromstring(data)
                for node in list(root):
                    if node.attrib.get("PartName") == "/docProps/custom.xml":
                        root.remove(node)
                data = ElementTree.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                )
            elif info.filename == "_rels/.rels":
                root = ElementTree.fromstring(data)
                for node in list(root):
                    if node.attrib.get("Type", "").endswith("/custom-properties"):
                        root.remove(node)
                data = ElementTree.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                )
            rewritten.writestr(info, data)
    return output.getvalue()


def _render_resume_docx(request: ResumeArtifactRequest) -> bytes:
    # compact_reference_guide + customer_pack-style left-aligned opening.
    document = _configure_docx(compact=True)
    _docx_title(
        document,
        request.student_name,
        "个人简历 · 用户确认信息的确定性排版版本",
    )
    _docx_key_value_table(
        document,
        [
            ("院系", request.dept),
            ("教育背景", request.education),
            ("邮箱", request.email),
            ("电话", request.phone),
        ],
    )
    if request.research_interests:
        document.add_heading("研究兴趣", level=1)
        document.add_paragraph("、".join(request.research_interests))
    if request.projects:
        document.add_heading("项目经历", level=1)
        for project in request.projects:
            document.add_heading(project.name, level=2)
            document.add_paragraph(project.detail or "未填写项目说明")
    if request.awards:
        document.add_heading("奖项与荣誉", level=1)
        _docx_key_value_table(
            document,
            [(str(index), item) for index, item in enumerate(request.awards, 1)],
        )
    if request.positions:
        document.add_heading("任职与实践", level=1)
        _docx_key_value_table(
            document,
            [(str(index), item) for index, item in enumerate(request.positions, 1)],
        )
    if request.target_advisor:
        document.add_heading("意向说明", level=1)
        _docx_callout(
            document,
            f"用户填写的意向对象：{request.target_advisor}。该名称未由系统核验，"
            "不代表对方正在招生或已同意接收材料。",
            warning=True,
        )
    document.add_paragraph()
    _docx_callout(
        document,
        "隐私与真实性提示：本文档只排版用户本次明确提供的信息，"
        "系统未核验经历真实性，也未向任何导师、邮箱或第三方发送文件。"
        "请在每次对外使用前逐项复核。",
        warning=True,
    )
    output = io.BytesIO()
    document.save(output)
    return _scrub_docx_metadata(output.getvalue())


def _report_plain_text(
    profile: StudentPortrait,
    outcome: MatchApplicationOutcome,
) -> str:
    lines = [
        "Tsing-RADAR 匹配报告",
        outcome.message,
        *[f"{label}：{value}" for label, value in _portrait_rows(profile)],
    ]
    for index, item in enumerate(outcome.items, 1):
        lines.append(
            f"{index}. {item['name']} 保守排序 {item['score']:.1f} "
            f"适配 {item['fit_score']:.1f} 覆盖 {item['evidence_coverage']:.0%} "
            f"置信 {item['evidence_confidence']:.0%}"
        )
        explanation = item.get("explanation") or {}
        lines.extend(
            f"支持：{claim['statement']}"
            for claim in explanation.get("supporting_evidence", [])
        )
        lines.extend(
            f"反证：{claim['statement']}"
            for claim in explanation.get("counter_evidence", [])
        )
        lines.extend(
            f"不确定性：{value}"
            for value in explanation.get("uncertainties", [])
        )
    return "\n".join(lines)


def _pdf_radar_section(outcome: MatchApplicationOutcome, styles: dict) -> list:
    """匹配报告的雷达图节：仅为拥有已审核六维评分的候选绘制（诚实空态兜底）。"""
    if outcome.status != "matched" or not outcome.items:
        return []
    from app.services.mentor_score_governance import public_score_bundles
    from app.services.radar_chart import (
        build_radar_series_for_advisor,
        render_radar_drawing,
    )

    try:
        bundles, status = public_score_bundles()
    except Exception:  # noqa: BLE001 —— 评分数据异常时按无数据呈现，不影响报告生成
        bundles, status = {}, {}
    font = _ensure_pdf_font()
    flowables = [_pdf_paragraph("导师特质雷达图（已审核评分）", styles["h1"])]
    drawn = 0
    for item in outcome.items[:3]:
        advisor_id = str(item.get("advisor_id") or "")
        if not advisor_id:
            continue
        series = build_radar_series_for_advisor(advisor_id, bundles)
        if series is None:
            continue
        release_version = status.get("release_version")
        sample_note = (
            f"样本来源：已审核评分发布 v{release_version}"
            if release_version
            else "样本来源：已审核评分发布"
        )
        block = [
            _pdf_paragraph(f"{item.get('name', advisor_id)}", styles["h2"]),
            render_radar_drawing(
                series=[series],
                title="六维特质（满分 100）",
                sample_note=sample_note,
                font_name=font,
            ),
            Spacer(1, 10),
        ]
        flowables.append(KeepTogether(block))
        drawn += 1
    if drawn == 0:
        flowables.append(
            _pdf_callout(
                "当前发布评分未覆盖这些候选导师，因此本报告不绘制雷达图；"
                "评分经审核发布后会自动出现在后续报告中。",
                styles,
                warning=True,
            )
        )
    return flowables


def _render_match_report_pdf(
    profile: StudentPortrait,
    outcome: MatchApplicationOutcome,
) -> bytes:
    styles = _pdf_styles()
    generated_at = _now()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Tsing-RADAR 匹配报告",
        author="Tsing-RADAR",
    )
    story = [
        _pdf_paragraph("Tsing-RADAR 匹配报告", styles["title"]),
        _pdf_paragraph(
            f"生成时间：{generated_at.astimezone().strftime('%Y-%m-%d %H:%M %Z')} · "
            "仅基于已确认画像与通过发布审核的数据",
            styles["subtitle"],
        ),
        _pdf_callout(
            outcome.message,
            styles,
            warning=outcome.status != "matched",
        ),
        _pdf_paragraph("已确认学生画像", styles["h1"]),
        _pdf_key_value_table(_portrait_rows(profile), styles),
    ]
    if outcome.status == "matched":
        story.append(_pdf_paragraph("证据化候选", styles["h1"]))
        for index, item in enumerate(outcome.items, 1):
            explanation = item.get("explanation") or {}
            block = [
                _pdf_paragraph(f"{index}. {item['name']}", styles["h2"]),
                _pdf_key_value_table(
                    [
                        ("保守排序分", f"{item['score']:.1f} / 100"),
                        ("适配分", f"{item['fit_score']:.1f} / 100"),
                        ("证据覆盖", f"{item['evidence_coverage']:.0%}"),
                        ("证据置信度", f"{item['evidence_confidence']:.0%}"),
                    ],
                    styles,
                ),
            ]
            for claim in explanation.get("supporting_evidence", []):
                block.append(
                    _pdf_paragraph(f"支持证据：{claim['statement']}", styles["body"])
                )
                for citation in claim.get("citations", []):
                    label = citation.get("citation") or citation.get("evidence_id")
                    source_url = citation.get("source_url")
                    block.append(
                        _pdf_paragraph(
                            f"来源：{label}"
                            + (f" · {source_url}" if source_url else ""),
                            styles["small"],
                        )
                    )
            for claim in explanation.get("counter_evidence", []):
                block.append(
                    _pdf_paragraph(f"反证/不利信号：{claim['statement']}", styles["body"])
                )
            for uncertainty in explanation.get("uncertainties", []):
                block.append(
                    _pdf_paragraph(f"不确定性：{uncertainty}", styles["body"])
                )
            for question in explanation.get("questions_to_verify", []):
                block.append(_pdf_paragraph(f"待核实：{question}", styles["body"]))
            story.append(KeepTogether(block))
    else:
        story.extend(
            [
                _pdf_paragraph("推荐结果", styles["h1"]),
                _pdf_callout(
                    "本报告不包含导师推荐名单。2027 官方招生目录只证明目录中的"
                    "院系、专业、方向、导师/导师组和备注等快照事实，"
                    "不能替代导师个人画像或推荐准确率。",
                    styles,
                    warning=True,
                ),
            ]
        )
    story.extend(_pdf_radar_section(outcome, styles))
    story.extend(
        [
            _pdf_paragraph("方法与边界", styles["h1"]),
            _pdf_key_value_table(
                [
                    (
                        "匹配状态",
                        _MATCH_STATUS_LABELS.get(outcome.status, outcome.status),
                    ),
                    (
                        "排序方法",
                        _MATCH_METHOD_LABELS.get(
                            str(
                                (outcome.meta.get("matching") or {}).get(
                                    "method_version",
                                    "未执行",
                                )
                            ),
                            str(
                                (outcome.meta.get("matching") or {}).get(
                                    "method_version",
                                    "未执行",
                                )
                            ),
                        ),
                    ),
                    (
                        "召回模式",
                        _RETRIEVAL_MODE_LABELS.get(
                            str(
                                (outcome.meta.get("matching") or {}).get(
                                    "retrieval_mode",
                                    "未执行",
                                )
                            ),
                            str(
                                (outcome.meta.get("matching") or {}).get(
                                    "retrieval_mode",
                                    "未执行",
                                )
                            ),
                        ),
                    ),
                    (
                        "发布记录数",
                        str(outcome.meta.get("published_records", 0)),
                    ),
                ],
                styles,
            ),
            Spacer(1, 10),
            _pdf_callout(
                "本报告不是录取、招生名额或合作承诺。缺失证据会降低覆盖与保守排序；"
                "任何待核实问题都应在采取行动前向公开来源或经授权主体确认。",
                styles,
                warning=True,
            ),
        ]
    )
    document.build(
        story,
        onFirstPage=_pdf_header_footer,
        onLaterPages=_pdf_header_footer,
    )
    return output.getvalue()


def _render_match_report_docx(
    profile: StudentPortrait,
    outcome: MatchApplicationOutcome,
) -> bytes:
    # standard_business_brief + memo_masthead opening.
    document = _configure_docx(compact=False)
    _docx_title(
        document,
        "Tsing-RADAR 匹配报告",
        "基于已确认画像与通过发布审核数据的证据化结果",
    )
    _docx_key_value_table(
        document,
        [
            ("生成时间", _now().astimezone().strftime("%Y-%m-%d %H:%M %Z")),
            (
                "匹配状态",
                _MATCH_STATUS_LABELS.get(outcome.status, outcome.status),
            ),
            (
                "发布记录数",
                str(outcome.meta.get("published_records", 0)),
            ),
            (
                "方法版本",
                _MATCH_METHOD_LABELS.get(
                    str(
                        (outcome.meta.get("matching") or {}).get(
                            "method_version",
                            "未执行",
                        )
                    ),
                    str(
                        (outcome.meta.get("matching") or {}).get(
                            "method_version",
                            "未执行",
                        )
                    ),
                ),
            ),
        ],
    )
    document.add_paragraph()
    _docx_callout(
        document,
        outcome.message,
        warning=outcome.status != "matched",
    )
    document.add_heading("已确认学生画像", level=1)
    _docx_key_value_table(document, _portrait_rows(profile))

    if outcome.status == "matched":
        document.add_heading("证据化候选", level=1)
        for index, item in enumerate(outcome.items, 1):
            document.add_heading(f"{index}. {item['name']}", level=2)
            _docx_key_value_table(
                document,
                [
                    ("保守排序分", f"{item['score']:.1f} / 100"),
                    ("适配分", f"{item['fit_score']:.1f} / 100"),
                    ("证据覆盖", f"{item['evidence_coverage']:.0%}"),
                    ("证据置信度", f"{item['evidence_confidence']:.0%}"),
                ],
            )
            explanation = item.get("explanation") or {}
            for claim in explanation.get("supporting_evidence", []):
                paragraph = document.add_paragraph()
                label = paragraph.add_run("支持证据：")
                _set_run_font(label, bold=True, color=_DOCX_DARK_BLUE)
                _set_run_font(paragraph.add_run(claim["statement"]))
                for citation in claim.get("citations", []):
                    source = document.add_paragraph()
                    source.paragraph_format.space_before = Pt(4)
                    source.paragraph_format.space_after = Pt(4)
                    value = citation.get("citation") or citation.get("evidence_id")
                    if citation.get("source_url"):
                        value += f" · {citation['source_url']}"
                    _set_run_font(source.add_run(f"来源：{value}"), size=9, color=_DOCX_MUTED)
            for claim in explanation.get("counter_evidence", []):
                _docx_callout(
                    document,
                    f"反证/不利信号：{claim['statement']}",
                    warning=True,
                )
            for uncertainty in explanation.get("uncertainties", []):
                document.add_paragraph(f"不确定性：{uncertainty}")
            for question in explanation.get("questions_to_verify", []):
                document.add_paragraph(f"待核实：{question}")
    else:
        document.add_heading("推荐结果", level=1)
        _docx_callout(
            document,
            "本报告不包含导师推荐名单。2027 官方招生目录只证明目录中的"
            "院系、专业、方向、导师/导师组和备注等快照事实，"
            "不能替代导师个人画像或推荐准确率。",
            warning=True,
        )

    document.add_heading("方法与边界", level=1)
    _docx_callout(
        document,
        "本报告不是录取、招生名额或合作承诺。缺失证据会降低覆盖与保守排序；"
        "任何待核实问题都应在采取行动前向公开来源或经授权主体确认。",
        warning=True,
    )
    output = io.BytesIO()
    document.save(output)
    return _scrub_docx_metadata(output.getvalue())


def create_resume_artifact(
    db: Session,
    *,
    owner_subject_id: str,
    channel: str,
    request: ResumeArtifactRequest,
    idempotency_key: str,
) -> PrivateDocument:
    if request.confirm_generation is not True:
        raise HTTPException(
            status_code=422,
            detail="生成前必须确认个人信息将写入私有文档",
        )
    operation = f"generate_resume:{channel}"
    claim = begin_idempotency(
        db,
        owner_subject_id=owner_subject_id,
        operation=operation,
        key=idempotency_key,
        payload=request.model_dump(mode="json"),
    )
    if claim.replayed:
        document = db.get(PrivateDocument, claim.record.resource_id)
        if document is None or document.owner_subject_id != owner_subject_id:
            raise HTTPException(status_code=410, detail="此前生成的私有文档已不可用")
        return document

    document: PrivateDocument | None = None
    phase = "render"
    try:
        if request.format == "pdf":
            payload = _render_resume_pdf(request)
            media_type = "application/pdf"
            filename = f"{request.student_name}-个人简历.pdf"
        else:
            payload = _render_resume_docx(request)
            media_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            filename = f"{request.student_name}-个人简历.docx"
        phase = "validate"
        safe_name, _extension, extracted, scan = validate_and_extract_document(
            payload=payload,
            filename=filename,
            media_type=media_type,
        )
        document = store_private_artifact(
            db,
            owner_subject_id=owner_subject_id,
            original_name=safe_name,
            payload=payload,
            media_type=media_type,
            document_kind="resume",
            extracted_text=extracted or _resume_plain_text(request),
            scan_result=scan,
            generation_context={
                "generator": "deterministic-local-template-v1",
                "format": request.format,
                "channel": channel,
                "external_model_used": False,
            },
            user_confirmed_at=_now(),
            commit=False,
        )
        add_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id,
            event_type="scan_completed",
            outcome="success",
            reason_code="scan_clean",
            scan_method=scan.method,
        )
        add_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id,
            event_type="generate_completed",
            outcome="success",
            reason_code="deterministic_private_document_generated",
            scan_method=scan.method,
        )
        complete_idempotency(
            db,
            record=claim.record,
            attempt_token=claim.attempt_token,
            resource_type="private_document",
            resource_id=document.document_id,
            commit=False,
        )
        db.commit()
        db.refresh(document)
        return document
    except Exception as exc:
        db.rollback()
        discard_private_artifact_object(document)
        fail_idempotency(
            db,
            record_id=claim.record.idempotency_id,
            attempt_token=claim.attempt_token,
            exc=exc,
            commit=False,
        )
        if phase == "validate" and isinstance(exc, HTTPException):
            event_type, reason_code = validation_reason(exc)
        else:
            event_type, reason_code = "generate_failed", "generation_failed"
        commit_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id if document else None,
            event_type=event_type,
            outcome="failed",
            reason_code=reason_code,
        )
        if isinstance(exc, HTTPException):
            raise
        raise HTTPException(
            status_code=503,
            detail=UNEXPECTED_FAILURE_DETAIL,
        ) from exc


def create_match_report_artifact(
    db: Session,
    *,
    owner_subject_id: str,
    channel: str,
    session_id: str,
    output_format: str,
    confirmed: bool,
    idempotency_key: str,
) -> tuple[PrivateDocument, MatchApplicationOutcome]:
    if confirmed is not True:
        raise HTTPException(
            status_code=422,
            detail="生成前必须确认画像与匹配结果将写入私有报告",
        )
    session = get_session(db, session_id, owner_subject_id)
    profile = confirmed_portrait(
        db,
        session_id=session_id,
        student_id=owner_subject_id,
    )
    outcome = run_confirmed_match(
        db,
        session_id=session_id,
        student_id=owner_subject_id,
    )
    operation = f"generate_match_report:{channel}"
    claim = begin_idempotency(
        db,
        owner_subject_id=owner_subject_id,
        operation=operation,
        key=idempotency_key,
        payload={
            "session_id": session_id,
            "format": output_format,
            "confirm_generation": confirmed,
        },
    )
    if claim.replayed:
        document = db.get(PrivateDocument, claim.record.resource_id)
        if document is None or document.owner_subject_id != owner_subject_id:
            raise HTTPException(status_code=410, detail="此前生成的匹配报告已不可用")
        return document, outcome
    document: PrivateDocument | None = None
    phase = "render"
    try:
        if output_format == "pdf":
            payload = _render_match_report_pdf(profile, outcome)
            media_type = "application/pdf"
            extension = "pdf"
        elif output_format == "docx":
            payload = _render_match_report_docx(profile, outcome)
            media_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            extension = "docx"
        else:
            raise HTTPException(status_code=422, detail="报告格式仅支持 pdf 或 docx")
        filename = (
            "Tsing-RADAR-匹配报告-"
            f"{_now().astimezone().strftime('%Y%m%d')}.{extension}"
        )
        phase = "validate"
        safe_name, _extension, extracted, scan = validate_and_extract_document(
            payload=payload,
            filename=filename,
            media_type=media_type,
        )
        document = store_private_artifact(
            db,
            owner_subject_id=owner_subject_id,
            original_name=safe_name,
            payload=payload,
            media_type=media_type,
            document_kind="match_report",
            extracted_text=extracted or _report_plain_text(profile, outcome),
            scan_result=scan,
            source_session_id=session_id,
            generation_context={
                "generator": "deterministic-evidence-report-v1",
                "format": output_format,
                "channel": channel,
                "outcome_status": outcome.status,
                "profile_version": int(session.profile_version or 1),
                "published_records": int(outcome.meta.get("published_records", 0)),
                "external_model_used": False,
            },
            user_confirmed_at=_now(),
            commit=False,
        )
        add_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id,
            event_type="scan_completed",
            outcome="success",
            reason_code="scan_clean",
            scan_method=scan.method,
        )
        add_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id,
            event_type="generate_completed",
            outcome="success",
            reason_code="deterministic_match_report_generated",
            scan_method=scan.method,
        )
        complete_idempotency(
            db,
            record=claim.record,
            attempt_token=claim.attempt_token,
            resource_type="private_document",
            resource_id=document.document_id,
            commit=False,
        )
        db.commit()
        db.refresh(document)
        return document, outcome
    except Exception as exc:
        db.rollback()
        discard_private_artifact_object(document)
        fail_idempotency(
            db,
            record_id=claim.record.idempotency_id,
            attempt_token=claim.attempt_token,
            exc=exc,
            commit=False,
        )
        if phase == "validate" and isinstance(exc, HTTPException):
            event_type, reason_code = validation_reason(exc)
        else:
            event_type, reason_code = "generate_failed", "generation_failed"
        commit_artifact_event(
            db,
            owner_subject_id=owner_subject_id,
            operation=operation,
            idempotency_key_digest=claim.record.key_digest,
            document_id=document.document_id if document else None,
            event_type=event_type,
            outcome="failed",
            reason_code=reason_code,
        )
        if isinstance(exc, HTTPException):
            raise
        raise HTTPException(
            status_code=503,
            detail=UNEXPECTED_FAILURE_DETAIL,
        ) from exc
